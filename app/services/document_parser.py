"""다중 포맷 문서 파서 (PDF / DOCX).

반환 타입: ParsedDocument
  - pages  : 페이지/섹션 단위 블록 목록
  - tables : 추출된 테이블 목록
  - meta   : 파일명·총 페이지 수 등 메타데이터

청킹: RecursiveCharacterTextSplitter (LangChain)
  - 한·중 문장 경계(.  。 ? !) 우선 분할
  - 토큰 수 근사: CJK 1자 ≈ 1토큰, 영문 4자 ≈ 1토큰
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# 문서 유형
# ---------------------------------------------------------------------------

class DocType(str, Enum):
    REGULATION = "규정"
    NOTICE     = "공지"
    GUIDE      = "안내"


# ---------------------------------------------------------------------------
# 토큰 수 근사 함수
# ---------------------------------------------------------------------------

# CJK Unified Ideographs + 한글 범위
_CJK_RE = re.compile(
    r"[\u1100-\u11FF"   # 한글 자모
    r"\uAC00-\uD7AF"    # 한글 음절
    r"\u4E00-\u9FFF"    # CJK 통합 한자
    r"\u3400-\u4DBF"    # CJK 확장 A
    r"\uF900-\uFAFF"    # CJK 호환 한자
    r"]"
)


def _approx_tokens(text: str) -> int:
    """CJK 문자 1자 ≈ 1토큰, 그 외 4문자 ≈ 1토큰으로 근사."""
    cjk_count = len(_CJK_RE.findall(text))
    other_count = len(text) - cjk_count
    return cjk_count + max(other_count // 4, 0)


# ---------------------------------------------------------------------------
# 청킹 설정
# ---------------------------------------------------------------------------

# 한·중 문장 경계를 우선 분할 기준으로 사용
_KO_SEPARATORS = [
    "\n\n",   # 단락 구분 (최우선)
    "\n",     # 줄바꿈
    "。",     # 중국어/일본어 마침표
    ". ",     # 영문 마침표 + 공백
    "? ",     # 의문문
    "! ",     # 감탄문
    ".",      # 마침표 (공백 없는 경우)
    "?",
    "!",
    " ",      # 단어 경계 (최후 수단)
]


# ---------------------------------------------------------------------------
# 데이터 모델
# ---------------------------------------------------------------------------

@dataclass
class PageBlock:
    """페이지(PDF) 또는 섹션(DOCX) 단위 텍스트 블록."""
    page_number: int          # PDF: 1-based 페이지 번호 / DOCX: 섹션 순번
    section: str              # 직전 헤딩 텍스트 (없으면 빈 문자열)
    text: str                 # 정제된 본문


@dataclass
class TableBlock:
    """추출된 테이블."""
    page_number: int
    section: str
    rows: list[list[str]]     # 행 × 열 문자열 그리드


@dataclass
class ParsedDocument:
    filename: str
    total_pages: int
    pages: list[PageBlock] = field(default_factory=list)
    tables: list[TableBlock] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        """모든 페이지 텍스트를 하나로 합친 문자열."""
        return "\n\n".join(p.text for p in self.pages if p.text)


# ---------------------------------------------------------------------------
# 텍스트 정제 유틸리티
# ---------------------------------------------------------------------------

# 머리글/바닥글 패턴: 페이지 번호, '- N -', 'Page N of M' 등
_HEADER_FOOTER_RE = re.compile(
    r"^(?:"
    r"\s*[-–—]\s*\d+\s*[-–—]\s*"      # - 1 -
    r"|Page\s+\d+\s+of\s+\d+"          # Page 1 of 5
    r"|\d+\s*/\s*\d+"                  # 1/5
    r"|={3,}"                           # ===...
    r")\s*$",
    re.IGNORECASE | re.MULTILINE,
)

# 연속 공백 정리 (한·중 유니코드 보존)
_MULTI_SPACE_RE  = re.compile(r"[ \t]+")
_MULTI_NEWLINE_RE = re.compile(r"\n{3,}")


def _clean(text: str) -> str:
    """공백 정규화, 머리글/바닥글 제거, 한·중 특수문자 보존."""
    # NFC 정규화 (한국어 자모 분리 방지)
    text = unicodedata.normalize("NFC", text)
    # 머리글/바닥글 제거
    text = _HEADER_FOOTER_RE.sub("", text)
    # 탭·연속 스페이스 → 단일 공백
    text = _MULTI_SPACE_RE.sub(" ", text)
    # 3줄 이상 빈 줄 → 2줄
    text = _MULTI_NEWLINE_RE.sub("\n\n", text)
    return text.strip()


def _is_heading_style(style_name: str) -> bool:
    return style_name.lower().startswith("heading")


# ---------------------------------------------------------------------------
# PDF 파서
# ---------------------------------------------------------------------------

class _PDFParser:
    # 페이지 상하 10 % 영역을 머리글/바닥글로 간주
    _HEADER_RATIO = 0.10
    _FOOTER_RATIO = 0.90

    def parse(self, file_path: str) -> ParsedDocument:
        doc = fitz.open(file_path)
        filename = Path(file_path).name
        result = ParsedDocument(filename=filename, total_pages=doc.page_count)
        current_section = ""

        for page in doc:
            page_num = page.number + 1
            body_text, tables = self._extract_page(page)
            body_text = _clean(body_text)

            # 섹션 힌트: 볼드 + 짧은 줄을 헤딩으로 추정
            for line in body_text.splitlines():
                stripped = line.strip()
                if stripped and len(stripped) < 60 and not stripped.endswith("."):
                    current_section = stripped
                    break

            if body_text:
                result.pages.append(PageBlock(
                    page_number=page_num,
                    section=current_section,
                    text=body_text,
                ))

            for tbl in tables:
                result.tables.append(TableBlock(
                    page_number=page_num,
                    section=current_section,
                    rows=tbl,
                ))

        doc.close()
        return result

    def _extract_page(self, page: fitz.Page) -> tuple[str, list[list[list[str]]]]:
        """본문 텍스트와 테이블을 분리 추출."""
        height = page.rect.height
        header_y = height * self._HEADER_RATIO
        footer_y = height * self._FOOTER_RATIO

        # 블록 단위 추출 (dict 모드)
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

        body_lines: list[str] = []
        tables: list[list[list[str]]] = []

        # 테이블 영역 감지 (PyMuPDF 1.23+)
        table_areas: list[fitz.Rect] = []
        try:
            for tbl in page.find_tables().tables:
                table_areas.append(tbl.bbox)
                rows = [
                    [_clean(cell or "") for cell in row]
                    for row in tbl.extract()
                ]
                tables.append(rows)
        except AttributeError:
            pass  # PyMuPDF 버전이 낮으면 테이블 감지 생략

        def _in_table(rect: fitz.Rect) -> bool:
            return any(rect.intersects(ta) for ta in table_areas)

        for block in blocks:
            if block.get("type") != 0:  # 0 = text block
                continue
            rect = fitz.Rect(block["bbox"])
            # 머리글/바닥글 영역 제외
            if rect.y1 < header_y or rect.y0 > footer_y:
                continue
            # 테이블 영역 제외 (텍스트 중복 방지)
            if _in_table(rect):
                continue

            for line in block.get("lines", []):
                line_text = "".join(
                    span["text"] for span in line.get("spans", [])
                )
                body_lines.append(line_text)

        return "\n".join(body_lines), tables


# ---------------------------------------------------------------------------
# DOCX 파서
# ---------------------------------------------------------------------------

class _DOCXParser:
    def parse(self, file_path: str) -> ParsedDocument:
        doc = DocxDocument(file_path)
        filename = Path(file_path).name

        # DOCX는 페이지 개념이 없으므로 섹션을 페이지 번호 대용으로 사용
        result = ParsedDocument(filename=filename, total_pages=0)
        section_index = 0
        current_section = ""
        buffer: list[str] = []

        def _flush(section: str, idx: int) -> None:
            text = _clean("\n".join(buffer))
            if text:
                result.pages.append(PageBlock(
                    page_number=idx,
                    section=section,
                    text=text,
                ))
            buffer.clear()

        body = doc.element.body
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                para_text = "".join(n.text or "" for n in child.iter(qn("w:t")))
                if not para_text.strip():
                    continue

                # 스타일 이름 추출
                style_elem = child.find(f".//{qn('w:pStyle')}")
                style_name = style_elem.get(qn("w:val"), "") if style_elem is not None else ""

                if _is_heading_style(style_name):
                    _flush(current_section, section_index)
                    section_index += 1
                    current_section = para_text.strip()
                else:
                    buffer.append(para_text)

            elif tag == "tbl":
                _flush(current_section, section_index)
                rows = self._extract_table(child)
                if rows:
                    result.tables.append(TableBlock(
                        page_number=section_index,
                        section=current_section,
                        rows=rows,
                    ))

        _flush(current_section, section_index)
        result.total_pages = section_index + 1
        return result

    def _extract_table(self, tbl_elem) -> list[list[str]]:
        rows: list[list[str]] = []
        for tr in tbl_elem.findall(f".//{qn('w:tr')}"):
            cells: list[str] = []
            for tc in tr.findall(f".//{qn('w:tc')}"):
                cell_text = "".join(
                    n.text or "" for n in tc.iter(qn("w:t"))
                )
                cells.append(_clean(cell_text))
            if cells:
                rows.append(cells)
        return rows


# ---------------------------------------------------------------------------
# 퍼블릭 인터페이스
# ---------------------------------------------------------------------------

class DocumentParser:
    """PDF / DOCX 문서를 ParsedDocument로 파싱하는 통합 클래스."""

    _SUPPORTED = {".pdf", ".docx"}

    def __init__(self) -> None:
        self._pdf = _PDFParser()
        self._docx = _DOCXParser()

    def parse(self, file_path: str) -> ParsedDocument:
        """파일 경로를 받아 ParsedDocument를 반환한다.

        Raises:
            ValueError: 지원하지 않는 확장자
            FileNotFoundError: 파일이 존재하지 않는 경우
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._pdf.parse(file_path)
        elif suffix == ".docx":
            return self._docx.parse(file_path)
        else:
            raise ValueError(
                f"지원하지 않는 파일 형식입니다: '{suffix}'. "
                f"지원 형식: {', '.join(sorted(self._SUPPORTED))}"
            )

    def chunk(
        self,
        doc: ParsedDocument,
        doc_type: DocType = DocType.GUIDE,
        chunk_size: int = 800,
        overlap: int = 100,
    ) -> list[dict]:
        """ParsedDocument를 의미 단위 청크로 분할 (RecursiveCharacterTextSplitter).

        Args:
            doc:        파싱된 문서 객체
            doc_type:   문서 유형 (규정 / 공지 / 안내)
            chunk_size: 최대 토큰 수 (CJK 1자 ≈ 1토큰, 영문 4자 ≈ 1토큰)
            overlap:    청크 간 오버랩 토큰 수

        Returns:
            청크 딕셔너리 목록::

                {
                    "text": str,
                    "metadata": {
                        "source_doc":  str,   # 원본 문서명
                        "page_num":    int,   # 페이지/섹션 번호
                        "chunk_index": int,   # 전체 청크 내 순서 (0-based)
                        "doc_type":    str,   # 문서 유형
                    }
                }
        """
        splitter = RecursiveCharacterTextSplitter(
            separators=_KO_SEPARATORS,
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            length_function=_approx_tokens,
            keep_separator=True,
            is_separator_regex=False,
        )

        chunks: list[dict] = []
        chunk_index = 0

        for page in doc.pages:
            if not page.text.strip():
                continue

            splits = splitter.split_text(page.text)
            for split in splits:
                if not split.strip():
                    continue
                chunks.append({
                    "text": split,
                    "metadata": {
                        "source_doc":  doc.filename,
                        "page_num":    page.page_number,
                        "chunk_index": chunk_index,
                        "doc_type":    doc_type.value,
                    },
                })
                chunk_index += 1

        # 페이지 블록이 없는 경우 full_text 전체를 대상으로 청킹
        if not chunks and doc.full_text:
            splits = splitter.split_text(doc.full_text)
            for split in splits:
                if not split.strip():
                    continue
                chunks.append({
                    "text": split,
                    "metadata": {
                        "source_doc":  doc.filename,
                        "page_num":    0,
                        "chunk_index": chunk_index,
                        "doc_type":    doc_type.value,
                    },
                })
                chunk_index += 1

        return chunks


document_parser = DocumentParser()
